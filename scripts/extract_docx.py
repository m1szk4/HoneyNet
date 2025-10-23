#!/usr/bin/env python3
"""
Extract text content from Word .docx file for analysis
"""

from docx import Document
import sys

def extract_docx_content(filepath):
    """Extract all text content from a .docx file"""
    try:
        doc = Document(filepath)
        full_text = []

        # Extract main document text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        return '\n'.join(full_text)

    except Exception as e:
        return f"Error reading document: {str(e)}"

if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding='utf-8')

    filepath = r"c:\Users\Miszka\Documents\GitHub\HoneyNet\Honeynet do analizy ataków na urządzenia IoT_ Projekt i wnioski dla IDS.docx"

    content = extract_docx_content(filepath)
    print(content)
